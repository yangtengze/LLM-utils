from docx import Document
import numpy as np
class DocxLoader:
    def __init__(self):
        pass
    def load(self,file_path):
        doc = Document(file_path)
        data = []
        for para in doc.paragraphs:
            # print(para.text)
            data.append(para.text)
        # return doc.paragraphs
        return np.array(data)