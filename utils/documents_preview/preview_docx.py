import docx

def preview_docx(file_path):
    """预览DOCX文件"""
    try:
        doc = docx.Document(file_path)
        
        # 提取文本和基本格式
        html_content = "<div class='docx-preview'>"
        for para in doc.paragraphs:
        # for para in doc.paragraphs[:50]:  # 限制段落数
            style = para.style.name if para.style else "Normal"
            html_content += f"<p class='para-{style.lower()}'>{para.text}</p>"
        
        html_content += "</div>"
        
        return {
            'type': 'html',
            'content': html_content
        }
    except Exception as e:
        return {
            'type': 'text',
            'content': f"预览DOCX失败：{str(e)}"
        }
